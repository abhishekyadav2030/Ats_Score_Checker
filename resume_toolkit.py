import streamlit as st
import pdfplumber
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import string
from collections import Counter

# Download NLTK data
@st.cache_resource
def download_nltk_data():
    nltk.download('punkt_tab')
    nltk.download('stopwords')

download_nltk_data()

# ---------- Black & White CSS ----------
def set_bw_css():
    st.markdown("""
    <style>
    .stApp { background: #ffffff; }
    h1, h2, h3, p, li, span, label { color: #000000 !important; }
    .stButton button { background: #000000; color: white; border-radius: 20px; border: 1px solid #000000; }
    .stButton button:hover { background: #333333; }
    .stFileUploader, .stTextArea textarea { background: #f5f5f5; border: 1px solid #cccccc; border-radius: 10px; }
    div[data-testid="stMetricValue"] { color: #000000 !important; font-weight: bold; }
    .stProgress > div > div > div > div { background: #000000; }
    .streamlit-expanderHeader { background: #f0f0f0; border-radius: 10px; color: black; }
    </style>
    """, unsafe_allow_html=True)

# ---------- Common text extraction ----------
def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_resume_text(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()
    if file_type == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_type == "docx":
        return extract_text_from_docx(uploaded_file)
    else:
        st.error("❌ Only PDF or DOCX")
        return None

# ========== ATS SCORE MODULE ==========
stop_words = set(stopwords.words('english'))

def preprocess(text):
    text = text.lower()
    text = re.sub(f'[{re.escape(string.punctuation)}0-9]', ' ', text)
    tokens = word_tokenize(text)
    tokens = [t for t in tokens if t not in stop_words and len(t) > 2]
    return " ".join(tokens)

def extract_keywords_from_jd(jd_text, top_n=20):
    processed = preprocess(jd_text)
    words = processed.split()
    freq_dist = nltk.FreqDist(words)
    return [word for word, _ in freq_dist.most_common(top_n)]

def keyword_match_score(resume_text, jd_keywords):
    resume_lower = resume_text.lower()
    matched = [kw for kw in jd_keywords if kw in resume_lower]
    missing = [kw for kw in jd_keywords if kw not in resume_lower]
    score = (len(matched) / len(jd_keywords)) * 100 if jd_keywords else 0
    return score, matched, missing

def cosine_similarity_score(resume_text, jd_text):
    processed_resume = preprocess(resume_text)
    processed_jd = preprocess(jd_text)
    if not processed_resume or not processed_jd:
        return 0.0
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([processed_resume, processed_jd])
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
    return similarity[0][0] * 100

def compute_ats_score(resume_text, jd_text):
    jd_keywords = extract_keywords_from_jd(jd_text, top_n=20)
    kw_score, matched, missing = keyword_match_score(resume_text, jd_keywords)
    cos_score = cosine_similarity_score(resume_text, jd_text)
    final_score = (0.6 * kw_score) + (0.4 * cos_score)
    return final_score, kw_score, cos_score, matched, missing

def ats_score_tab():
    st.markdown("### 🎯 ATS Score")
    uploaded_file = st.file_uploader("📎 Resume (PDF/DOCX)", type=["pdf", "docx"], key="ats_upload")
    jd_text = st.text_area("📝 Job description", height=150, key="jd")
    
    if st.button("🔍 Check Score"):
        if not uploaded_file or not jd_text.strip():
            st.warning("⚠️ Upload resume and paste JD")
            return
        resume_text = extract_resume_text(uploaded_file)
        if not resume_text:
            st.error("❌ Could not read text")
            return
        
        final, kw, cos, matched, missing = compute_ats_score(resume_text, jd_text)
        
        st.metric("🏆 ATS Score", f"{final:.1f}%")
        st.progress(final / 100)
        col1, col2 = st.columns(2)
        col1.metric("🎯 Keyword match", f"{kw:.1f}%")
        col2.metric("🧠 Similarity", f"{cos:.1f}%")
        
        st.subheader("✅ Matched")
        st.write(", ".join(matched) if matched else "None")
        st.subheader("❌ Missing")
        st.write(", ".join(missing) if missing else "None")
        if missing:
            st.info("💡 Add missing keywords")

# ========== STRENGTH ANALYZER MODULE ==========
action_verbs = [
    "built", "created", "developed", "designed", "implemented", "led",
    "managed", "increased", "improved", "reduced", "solved", "optimised",
    "accelerated", "achieved", "negotiated", "presented", "analysed",
    "coordinated", "launched", "delivered", "trained", "mentored"
]

buzzwords = [
    "hardworking", "team player", "go getter", "synergy", "think outside the box",
    "results driven", "self motivated", "dynamic", "proactive", "detail oriented"
]

def count_action_verbs(text):
    text_lower = text.lower()
    found = [v for v in action_verbs if v in text_lower]
    return len(found), found

def count_numbers(text):
    numbers = re.findall(r'\b\d+(?:\.\d+)?%?\b', text)
    meaningful = [n for n in numbers if '%' in n or (n.isdigit() and int(n) > 10)]
    return len(meaningful)

def check_sections(text):
    return {
        "Experience": re.search(r'(?i)(work|professional|employment) experience', text) is not None,
        "Skills": re.search(r'(?i)skills|technical skills', text) is not None,
        "Education": re.search(r'(?i)education', text) is not None,
        "Projects": re.search(r'(?i)projects', text) is not None
    }

def count_buzzwords(text):
    text_lower = text.lower()
    found = [b for b in buzzwords if b in text_lower]
    return len(found), found

def readability_score(text):
    words = len(text.split())
    sentences = len(re.findall(r'[.!?]+', text))
    if sentences == 0:
        return "⚠️ No sentences"
    avg = words / sentences
    if avg < 12:
        return "👍 Good"
    elif avg < 20:
        return "👌 Okay"
    else:
        return "⚠️ Too long sentences"

def format_check(uploaded_file, text):
    issues = []
    ext = uploaded_file.name.split('.')[-1].lower()
    if ext == "pdf" and len(text.strip()) < 100:
        issues.append("❌ PDF may be scanned")
    elif ext == "docx":
        doc = Document(uploaded_file)
        if len(doc.tables) > 0:
            issues.append("⚠️ Tables detected")
        for para in doc.paragraphs:
            if 'shape' in para._element.xml:
                issues.append("⚠️ Images detected")
    return issues

def strength_tab():
    st.markdown("### ⚡ Resume Strength")
    uploaded_file = st.file_uploader("📎 Resume (PDF/DOCX)", type=["pdf", "docx"], key="strength_upload")
    
    if uploaded_file:
        text = extract_resume_text(uploaded_file)
        if not text:
            st.error("❌ Could not extract text")
            return
        
        verb_count, verb_list = count_action_verbs(text)
        st.metric("💪 Action Verbs", verb_count)
        if verb_list:
            st.write("✅", ", ".join(verb_list[:10]))
        
        num_count = count_numbers(text)
        st.metric("📊 Metrics (numbers/%)", num_count)
        if num_count == 0:
            st.warning("Add numbers like 'increased 20%'")
        
        sections = check_sections(text)
        st.subheader("📑 Sections")
        for sec, exists in sections.items():
            st.write("✅" if exists else "❌", sec)
        
        buzz_count, buzz_list = count_buzzwords(text)
        if buzz_count > 0:
            st.warning(f"⚠️ {buzz_count} buzzwords: {', '.join(buzz_list)}")
        
        read_grade = readability_score(text)
        st.info(f"📖 Readability: {read_grade}")
        
        format_issues = format_check(uploaded_file, text)
        if format_issues:
            for issue in format_issues:
                st.error(issue)
        
        # Overall strength score
        section_score = sum(sections.values()) / len(sections) * 100
        total = (verb_count * 5 + num_count * 10 + section_score) / (5 + 10 + 1)
        total = min(100, total)
        st.markdown("---")
        st.metric("🏆 Strength Score", f"{total:.0f}%")
        
        with st.expander("📄 Preview"):
            st.text(text[:800])

# ========== MAIN APP ==========
def main():
    st.set_page_config(page_title="Resume Toolkit", page_icon="📄")
    set_bw_css()
    st.markdown("<h1>📄 Resume Toolkit</h1>", unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["🎯 ATS Score", "⚡ Strength Analyzer"])
    
    with tab1:
        ats_score_tab()
    with tab2:
        strength_tab()
    
    st.markdown("---")
    st.caption("Black & White • Clean • Recruiter‑ready")

if __name__ == "__main__":
    main()